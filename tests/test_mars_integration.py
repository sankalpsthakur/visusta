"""
MARS backend integration tests — Phase 5.

These tests exercise multi-step workflows across the full stack:
real SQLite DB (migrations applied), real routers, real mars logic.

All tests skip gracefully under the system Python (no fastapi).
"""
from __future__ import annotations

import time
from pathlib import Path

import pytest

pytest.importorskip("fastapi")

from fastapi.testclient import TestClient  # noqa: E402
from docx import Document  # noqa: E402

import db.connection as db_connection_module  # noqa: E402


# ── Fixtures ───────────────────────────────────────────────────────────────────

@pytest.fixture()
def test_db_path(tmp_path: Path) -> Path:
    return tmp_path / "test_mars_integration.db"


@pytest.fixture()
def client(test_db_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    monkeypatch.setattr(db_connection_module, "DB_PATH", test_db_path)
    from api.main import app
    with TestClient(app, raise_server_exceptions=True) as c:
        yield c


# ── Helpers ────────────────────────────────────────────────────────────────────

CLIENT = "gerold-foods"


def _create_template(client: TestClient, name: str = "Integration Template") -> dict:
    return client.post("/api/templates", json={"name": name, "base_locale": "en"}).json()


def _create_draft(
    client: TestClient,
    title: str = "Integration Draft",
    locale: str = "en",
    period: str = "2026-02",
) -> dict:
    resp = client.post(
        f"/api/clients/{CLIENT}/drafts",
        json={"title": title, "primary_locale": locale, "period": period},
    )
    assert resp.status_code == 201
    return resp.json()


def _compose(client: TestClient, draft_id: int) -> dict:
    """Kick off the async compose job and poll until the revision is materialized."""
    resp = client.post(f"/api/clients/{CLIENT}/drafts/{draft_id}/compose")
    assert resp.status_code == 202, resp.text
    body = resp.json()
    assert "job_id" in body
    job_id = body["job_id"]
    # Poll the status endpoint. BackgroundTask runs in the TestClient's threadpool
    # with the stub LLM, so "done" lands within a handful of iterations.
    for _ in range(50):
        status = client.get(
            f"/api/clients/{CLIENT}/drafts/{draft_id}/compose/{job_id}"
        )
        assert status.status_code == 200, status.text
        data = status.json()
        if data["status"] == "done":
            assert data["revision"] is not None, "compose job finished without revision"
            return data["revision"]
        if data["status"] == "failed":
            raise AssertionError(f"Compose job failed: {data.get('error')}")
        time.sleep(0.05)
    raise AssertionError(f"Compose job {job_id} did not complete in time")


def _translate(client: TestClient, draft_id: int, target_locale: str) -> dict:
    """Kick off the async translate job and poll until the revision is materialized."""
    resp = client.post(
        f"/api/clients/{CLIENT}/drafts/{draft_id}/translate",
        params={"target_locale": target_locale},
    )
    assert resp.status_code == 202, resp.text
    body = resp.json()
    assert "job_id" in body
    job_id = body["job_id"]
    for _ in range(50):
        status = client.get(
            f"/api/clients/{CLIENT}/drafts/{draft_id}/translate/{job_id}"
        )
        assert status.status_code == 200, status.text
        data = status.json()
        if data["status"] == "done":
            assert data["revision"] is not None, "translate job finished without revision"
            return data["revision"]
        if data["status"] == "failed":
            raise AssertionError(f"Translate job failed: {data.get('error')}")
        time.sleep(0.05)
    raise AssertionError(f"Translate job {job_id} did not complete in time")


def _send_chat(
    client: TestClient,
    draft_id: int,
    *,
    content: str,
    section_id: str,
    role: str = "user",
) -> dict:
    """Fire a section-scoped chat (202) and poll for the assistant message."""
    resp = client.post(
        f"/api/clients/{CLIENT}/drafts/{draft_id}/chat",
        json={"role": role, "content": content, "section_id": section_id},
    )
    assert resp.status_code == 202, resp.text
    body = resp.json()
    assert "job_id" in body
    job_id = body["job_id"]
    for _ in range(50):
        status = client.get(
            f"/api/clients/{CLIENT}/drafts/{draft_id}/chat/{job_id}"
        )
        assert status.status_code == 200, status.text
        data = status.json()
        if data["status"] == "done":
            assert data["message"] is not None, "chat job finished without message"
            return data["message"]
        if data["status"] == "failed":
            raise AssertionError(f"Chat job failed: {data.get('error')}")
        time.sleep(0.05)
    raise AssertionError(f"Chat job {job_id} did not complete in time")


def _transition(client: TestClient, draft_id: int, target: str) -> dict:
    resp = client.post(
        f"/api/clients/{CLIENT}/drafts/{draft_id}/transition",
        params={"target_status": target},
    )
    assert resp.status_code == 200
    return resp.json()


def _approve_draft(client: TestClient, draft_id: int) -> dict:
    """Drive a draft through composing→review→approval→approved."""
    revision = _compose(client, draft_id)
    _transition(client, draft_id, "review")
    _transition(client, draft_id, "approval")
    for section in revision["sections"]:
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft_id}/approve",
            json={"section_id": section["section_id"], "status": "approved", "reviewer": "integration"},
        )
        assert resp.status_code == 200
    detail = client.get(f"/api/clients/{CLIENT}/drafts/{draft_id}").json()
    assert detail["status"] == "approved"
    return detail


# ── Template CRUD → clone → version inheritance ────────────────────────────────

class TestTemplateCloneWorkflow:
    def test_clone_inherits_base_locale(self, client: TestClient) -> None:
        src = _create_template(client, "Source")
        clone = client.post(
            f"/api/templates/{src['id']}/clone",
            json={"name": "Clone A"},
        ).json()
        assert clone["base_locale"] == src["base_locale"]

    def test_clone_gets_independent_id(self, client: TestClient) -> None:
        src = _create_template(client, "Source B")
        clone = client.post(
            f"/api/templates/{src['id']}/clone",
            json={"name": "Clone B"},
        ).json()
        assert clone["id"] != src["id"]

    def test_new_version_on_clone_starts_at_version_1(self, client: TestClient) -> None:
        src = _create_template(client, "Source C")
        clone = client.post(
            f"/api/templates/{src['id']}/clone",
            json={"name": "Clone C"},
        ).json()
        resp = client.post(
            f"/api/templates/{clone['id']}/versions",
            json={"sections_json": [{"section_id": "s1", "heading": "Intro"}], "changelog_note": "v2"},
        )
        assert resp.status_code == 201
        # Clone of versionless template starts at version 0, so first version gives 1
        assert resp.json()["version_number"] == 1

    def test_version_numbering_increments_per_template(self, client: TestClient) -> None:
        tmpl = _create_template(client, "Version Check")
        client.post(f"/api/templates/{tmpl['id']}/versions", json={"changelog_note": "v2"})
        resp = client.post(f"/api/templates/{tmpl['id']}/versions", json={"changelog_note": "v3"})
        # New templates start at current_version=0 until an explicit first version is created.
        assert resp.json()["version_number"] == 2

    def test_versions_from_source_not_shared_with_clone(self, client: TestClient) -> None:
        src = _create_template(client, "Source D")
        # Add 2 extra versions to source
        client.post(f"/api/templates/{src['id']}/versions", json={})
        client.post(f"/api/templates/{src['id']}/versions", json={})
        # Clone is independent — should only have 1 version (the initial)
        clone = client.post(
            f"/api/templates/{src['id']}/clone",
            json={"name": "Clone D"},
        ).json()
        versions = client.get(f"/api/templates/{clone['id']}/versions").json()
        assert len(versions) == 1

    def test_template_full_crud_lifecycle(self, client: TestClient) -> None:
        # Create
        tmpl = _create_template(client, "CRUD Life")
        assert tmpl["name"] == "CRUD Life"
        # Update
        updated = client.put(
            f"/api/templates/{tmpl['id']}",
            json={"name": "CRUD Life v2", "base_locale": "fr"},
        ).json()
        assert updated["name"] == "CRUD Life v2"
        # Read
        fetched = client.get(f"/api/templates/{tmpl['id']}").json()
        assert fetched["name"] == "CRUD Life v2"
        # Delete
        client.delete(f"/api/templates/{tmpl['id']}")
        assert client.get(f"/api/templates/{tmpl['id']}").status_code == 404


# ── Draft generation → revision history → chat ────────────────────────────────

class TestDraftRevisionHistoryWorkflow:
    def test_compose_then_compose_builds_history(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        _compose(client, draft["id"])
        revisions = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/revisions").json()
        assert len(revisions) == 2
        nums = [r["revision_number"] for r in revisions]
        assert sorted(nums) == [1, 2]

    def test_revision_numbers_are_sequential(self, client: TestClient) -> None:
        draft = _create_draft(client)
        for _ in range(3):
            _compose(client, draft["id"])
        revisions = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/revisions").json()
        assert [r["revision_number"] for r in revisions] == [1, 2, 3]

    def test_compose_updates_current_revision_id(self, client: TestClient) -> None:
        draft = _create_draft(client)
        rev1 = _compose(client, draft["id"])
        draft_detail = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}").json()
        assert draft_detail["current_revision_id"] == rev1["id"]
        rev2 = _compose(client, draft["id"])
        draft_detail = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}").json()
        assert draft_detail["current_revision_id"] == rev2["id"]

    def test_chat_message_linked_to_current_revision(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        msg = _send_chat(
            client,
            draft["id"],
            content="Please expand this section.",
            section_id="exec",
        )
        current = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}").json()
        assert msg["revision_id"] == current["current_revision_id"]

    def test_chat_messages_accumulate_across_composes(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "First message."},
        )
        _compose(client, draft["id"])
        client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "Second message."},
        )
        all_msgs = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/chat").json()
        assert len(all_msgs) == 4
        assert [message["role"] for message in all_msgs] == ["user", "assistant", "user", "assistant"]

    def test_get_specific_revision_by_id(self, client: TestClient) -> None:
        draft = _create_draft(client)
        rev1 = _compose(client, draft["id"])
        rev2 = _compose(client, draft["id"])
        # Fetch rev1 directly
        fetched = client.get(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/revisions/{rev1['id']}"
        ).json()
        assert fetched["id"] == rev1["id"]
        assert fetched["revision_number"] == 1
        assert fetched["id"] != rev2["id"]

    def test_sections_present_in_revision_detail(self, client: TestClient) -> None:
        draft = _create_draft(client)
        rev = _compose(client, draft["id"])
        assert "sections" in rev
        assert len(rev["sections"]) >= 1
        for sec in rev["sections"]:
            assert "section_id" in sec
            assert "heading" in sec
            assert "locale" in sec


# ── Approval workflow end-to-end ───────────────────────────────────────────────

@pytest.mark.usefixtures("stub_pdf_export")
class TestApprovalWorkflowEndToEnd:
    def test_full_state_machine_composing_to_approved(self, client: TestClient) -> None:
        draft = _create_draft(client)
        result = _approve_draft(client, draft["id"])
        assert result["status"] == "approved"

    def test_pdf_export_succeeds_after_approval(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _approve_draft(client, draft["id"])
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 201
        payload = resp.json()
        assert payload["format"] == "pdf"
        assert payload["status"] == "completed"

    def test_pdf_export_blocked_in_review_status(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        _transition(client, draft["id"], "review")
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 403

    def test_pdf_export_blocked_in_composing_status(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 403

    def test_pdf_export_allowed_after_exported_transition(self, client: TestClient) -> None:
        """exported status is also allowed (draft already exported once)."""
        draft = _create_draft(client)
        _approve_draft(client, draft["id"])
        _transition(client, draft["id"], "exported")
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 201
        assert resp.json()["status"] == "completed"

    def test_section_approvals_visible_after_approve(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/approve",
            json={"section_id": "executive_summary", "status": "approved", "reviewer": "alice"},
        )
        approvals = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/approvals").json()
        assert len(approvals) == 1
        assert approvals[0]["section_id"] == "executive_summary"
        assert approvals[0]["status"] == "approved"

    def test_multiple_section_approvals(self, client: TestClient) -> None:
        draft = _create_draft(client)
        revision = _compose(client, draft["id"])
        section_ids = [s["section_id"] for s in revision["sections"]]
        for sec_id in section_ids[:3]:
            client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/approve",
                json={"section_id": sec_id, "status": "approved", "reviewer": "carol"},
            )
        approvals = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/approvals").json()
        assert len(approvals) == len(section_ids[:3])
        approved_ids = {a["section_id"] for a in approvals}
        assert approved_ids == set(section_ids[:3])

    def test_approval_upsert_does_not_duplicate(self, client: TestClient) -> None:
        draft = _create_draft(client)
        revision = _compose(client, draft["id"])
        real_section_id = revision["sections"][0]["section_id"]
        for status in ("needs_revision", "approved"):
            client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/approve",
                json={"section_id": real_section_id, "status": status, "reviewer": "dave"},
            )
        approvals = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/approvals").json()
        matching = [a for a in approvals if a["section_id"] == real_section_id]
        assert len(matching) == 1
        assert matching[0]["status"] == "approved"

    def test_invalid_state_machine_transition_blocked(self, client: TestClient) -> None:
        draft = _create_draft(client)
        # composing → approved is not a valid direct transition
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/transition",
            params={"target_status": "approved"},
        )
        assert resp.status_code == 422

    def test_archived_draft_not_updatable(self, client: TestClient) -> None:
        draft = _create_draft(client)
        client.delete(f"/api/clients/{CLIENT}/drafts/{draft['id']}")
        # archived → anything is blocked
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/transition",
            params={"target_status": "composing"},
        )
        assert resp.status_code == 422


# ── DOCX round-trip ────────────────────────────────────────────────────────────

class TestDocxRoundTrip:
    def _make_docx(self, tmp_path: Path, headings: list[str]) -> Path:
        doc = Document()
        for heading in headings:
            doc.add_heading(heading, level=1)
            doc.add_paragraph(f"Content for {heading}.")
        path = tmp_path / "roundtrip.docx"
        doc.save(str(path))
        return path

    def test_import_docx_creates_revision(self, client: TestClient, tmp_path: Path) -> None:
        draft = _create_draft(client)
        docx_path = self._make_docx(tmp_path, ["Executive Summary", "Risk Factors"])
        with open(docx_path, "rb") as f:
            resp = client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports/import-docx",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        assert resp.status_code == 201
        data = resp.json()
        assert data["draft_id"] == draft["id"]
        assert data["sections_imported"] == 2

    def test_imported_sections_become_current_revision(self, client: TestClient, tmp_path: Path) -> None:
        draft = _create_draft(client)
        docx_path = self._make_docx(tmp_path, ["Introduction", "Conclusion"])
        with open(docx_path, "rb") as f:
            import_result = client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports/import-docx",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            ).json()
        detail = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}").json()
        assert detail["current_revision_id"] == import_result["revision_id"]

    def test_import_after_compose_adds_second_revision(self, client: TestClient, tmp_path: Path) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        docx_path = self._make_docx(tmp_path, ["Section A"])
        with open(docx_path, "rb") as f:
            client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports/import-docx",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        revisions = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/revisions").json()
        assert len(revisions) == 2

    def test_docx_export_job_created_for_unapproved_draft(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _compose(client, draft["id"])
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "docx", "locale": "en"},
        )
        assert resp.status_code == 201
        job = resp.json()
        assert job["format"] == "docx"
        assert job["status"] == "completed"

    def test_export_jobs_listed_after_creation(self, client: TestClient) -> None:
        draft = _create_draft(client)
        _approve_draft(client, draft["id"])
        client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        jobs = client.get(f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports").json()
        assert len(jobs) == 1
        assert jobs[0]["format"] == "pdf"

    def test_import_section_count_matches_docx_headings(self, client: TestClient, tmp_path: Path) -> None:
        headings = ["Section 1", "Section 2", "Section 3", "Section 4"]
        draft = _create_draft(client)
        docx_path = self._make_docx(tmp_path, headings)
        with open(docx_path, "rb") as f:
            result = client.post(
                f"/api/clients/{CLIENT}/drafts/{draft['id']}/exports/import-docx",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            ).json()
        assert result["sections_imported"] == len(headings)


# ── Source/keyword → proposal → approval ──────────────────────────────────────

class TestSourceKeywordWorkflow:
    def test_create_keywords_then_preview(self, client: TestClient) -> None:
        client.post(
            f"/api/clients/{CLIENT}/keywords",
            json={"phrase": "packaging regulation", "locale": "en"},
        )
        client.post(
            f"/api/clients/{CLIENT}/keywords",
            json={"phrase": "GHG reporting", "locale": "en"},
        )
        resp = client.post(
            f"/api/clients/{CLIENT}/keywords/preview",
            json={
                "phrases": ["packaging regulation", "GHG reporting"],
                "locale": "en",
                "sample_text": "The GHG reporting deadline for packaging regulation is Q3.",
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["match_count"] == 2

    def test_suggest_proposals_stores_to_db(self, client: TestClient) -> None:
        resp = client.post(f"/api/clients/{CLIENT}/source-proposals/suggest")
        assert resp.status_code == 201
        proposals = resp.json()
        # Proposals from stub LLM get stored; verify they appear in list
        listed = client.get(f"/api/clients/{CLIENT}/source-proposals").json()
        assert len(listed) == len(proposals)

    def test_approve_then_filter_by_status(self, client: TestClient) -> None:
        import db.connection as dbc
        with dbc.get_db() as conn:
            cur = conn.execute(
                "INSERT INTO source_proposals (client_id, url, title, proposed_by) VALUES (?,?,?,?)",
                (CLIENT, "https://eur-lex.europa.eu/a", "EUR-Lex A", "scout"),
            )
            pid1 = cur.lastrowid
            cur = conn.execute(
                "INSERT INTO source_proposals (client_id, url, title, proposed_by) VALUES (?,?,?,?)",
                (CLIENT, "https://eur-lex.europa.eu/b", "EUR-Lex B", "scout"),
            )
            pid2 = cur.lastrowid

        # Approve one
        client.post(
            f"/api/clients/{CLIENT}/source-proposals/{pid1}/action",
            json={"action": "approve", "reviewer": "alice"},
        )

        approved = client.get(
            f"/api/clients/{CLIENT}/source-proposals?status=approved"
        ).json()
        pending = client.get(
            f"/api/clients/{CLIENT}/source-proposals?status=pending"
        ).json()

        assert all(p["status"] == "approved" for p in approved)
        assert any(p["id"] == pid1 for p in approved)
        assert any(p["id"] == pid2 for p in pending)

    def test_reject_proposal_removes_from_pending(self, client: TestClient) -> None:
        import db.connection as dbc
        with dbc.get_db() as conn:
            cur = conn.execute(
                "INSERT INTO source_proposals (client_id, url, proposed_by) VALUES (?,?,?)",
                (CLIENT, "https://reject.me", "scout"),
            )
            pid = cur.lastrowid

        client.post(
            f"/api/clients/{CLIENT}/source-proposals/{pid}/action",
            json={"action": "reject", "reviewer": "bob", "note": "Not relevant"},
        )

        pending = client.get(
            f"/api/clients/{CLIENT}/source-proposals?status=pending"
        ).json()
        assert all(p["id"] != pid for p in pending)

    def test_keyword_soft_delete_hides_from_list(self, client: TestClient) -> None:
        kw = client.post(
            f"/api/clients/{CLIENT}/keywords",
            json={"phrase": "social compliance", "locale": "en"},
        ).json()
        client.delete(f"/api/clients/{CLIENT}/keywords/{kw['id']}")
        listed = client.get(f"/api/clients/{CLIENT}/keywords").json()
        assert all(k["id"] != kw["id"] for k in listed)

    def test_keyword_update_persists(self, client: TestClient) -> None:
        kw = client.post(
            f"/api/clients/{CLIENT}/keywords",
            json={"phrase": "water usage", "locale": "en", "weight": 1.0},
        ).json()
        updated = client.put(
            f"/api/clients/{CLIENT}/keywords/{kw['id']}",
            json={"phrase": "water footprint", "locale": "en", "weight": 4.0},
        ).json()
        assert updated["phrase"] == "water footprint"
        assert updated["weight"] == 4.0
        # Verify persistence
        fetched = client.get(f"/api/clients/{CLIENT}/keywords/{kw['id']}").json()
        assert fetched["phrase"] == "water footprint"


# ── Locale propagation workflow ────────────────────────────────────────────────

class TestLocaleWorkflow:
    def test_compose_preserves_primary_locale(self, client: TestClient) -> None:
        draft = _create_draft(client, locale="de")
        rev = _compose(client, draft["id"])
        for sec in rev["sections"]:
            assert sec["locale"] == "de"

    def test_translate_creates_new_revision_with_target_locale(self, client: TestClient) -> None:
        draft = _create_draft(client, locale="en")
        _compose(client, draft["id"])
        rev = _translate(client, draft["id"], "fr")
        for sec in rev["sections"]:
            assert sec["locale"] == "fr"
        headings = {sec["heading"] for sec in rev["sections"]}
        assert "Résumé exécutif" in headings or "Actions critiques" in headings

    def test_translate_increments_revision_number(self, client: TestClient) -> None:
        draft = _create_draft(client, locale="en")
        _compose(client, draft["id"])
        rev = _translate(client, draft["id"], "de")
        assert rev["revision_number"] == 2

    def test_locale_settings_roundtrip(self, client: TestClient) -> None:
        payload = {"primary_locale": "es", "enabled_locales": ["en", "de", "es", "fr"]}
        client.put(f"/api/clients/{CLIENT}/locale-settings", json=payload)
        fetched = client.get(f"/api/clients/{CLIENT}/locale-settings").json()
        assert fetched["primary_locale"] == "es"
        assert "es" in fetched["enabled_locales"]
        assert "fr" in fetched["enabled_locales"]

    def test_invalid_translate_locale_rejected(self, client: TestClient) -> None:
        draft = _create_draft(client, locale="en")
        _compose(client, draft["id"])
        resp = client.post(
            f"/api/clients/{CLIENT}/drafts/{draft['id']}/translate",
            params={"target_locale": "zz"},
        )
        assert resp.status_code == 422

    def test_all_24_locales_available(self, client: TestClient) -> None:
        # 24 EU + 2 Norwegian (nb, nn) seeded by migration 003
        locales = client.get("/api/locales").json()
        codes = {lc["code"] for lc in locales}
        assert len(codes) == 26
        # Spot-check key EU languages plus Norwegian variants
        for code in ("en", "de", "fr", "es", "it", "nl", "pl", "sv", "nb", "nn"):
            assert code in codes


# ── Multi-client isolation ─────────────────────────────────────────────────────

class TestClientIsolation:
    def test_drafts_isolated_per_client(self, client: TestClient) -> None:
        """Drafts created for gerold-foods are not visible to a different client."""
        _create_draft(client, title="GF Draft")
        # pf-emea is not a registered client — should 404
        resp = client.get("/api/clients/pf-emea/drafts")
        assert resp.status_code == 404

    def test_keywords_isolated_per_client(self, client: TestClient) -> None:
        client.post(
            f"/api/clients/{CLIENT}/keywords",
            json={"phrase": "isolated keyword", "locale": "en"},
        )
        # Different client returns 404
        resp = client.get("/api/clients/another-client-xyz/keywords")
        assert resp.status_code == 404

    def test_get_draft_wrong_client_404(self, client: TestClient) -> None:
        draft = _create_draft(client)
        resp = client.get(f"/api/clients/wrong-client/drafts/{draft['id']}")
        assert resp.status_code == 404
