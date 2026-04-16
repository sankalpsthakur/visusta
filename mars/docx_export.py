"""
DOCX export — structured generator for DraftSection data.

Handles: paragraph, heading, bullet_list, table block types.
Adds Key Facts and References sub-sections per section.
Supports optional client_branding (primary_color, company_name).
"""

from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import List

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from api.schemas_mars import DraftSection


def _parse_rgb(hex_color: str) -> RGBColor | None:
    """Parse '#RRGGBB' or 'RRGGBB' into RGBColor; return None on failure."""
    try:
        h = hex_color.lstrip("#")
        if len(h) != 6:
            return None
        r, g, b = int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)
        return RGBColor(r, g, b)
    except (ValueError, AttributeError):
        return None


def _apply_heading_color(paragraph, rgb: RGBColor) -> None:
    """Apply an RGBColor to every run in a heading paragraph."""
    for run in paragraph.runs:
        run.font.color.rgb = rgb


def _add_title_page(doc: Document, title: str, locale: str) -> None:
    """Insert a simple title page with title, locale, and generation date."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()  # spacer

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Locale: {locale}\nGenerated: {date.today().isoformat()}")

    doc.add_page_break()


def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(str(text))


def _add_numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(str(text))


def _citation_text(citation: object) -> str:
    """Render citations readably across legacy and dict-shaped revisions."""
    if isinstance(citation, str):
        return citation
    label = getattr(citation, "label", None)
    url = getattr(citation, "url", None)
    if label and url:
        return f"{label} - {url}"
    if label:
        return str(label)
    if isinstance(citation, dict):
        dict_label = citation.get("label")
        dict_url = citation.get("url")
        if dict_label and dict_url:
            return f"{dict_label} - {dict_url}"
        if dict_label:
            return str(dict_label)
    return str(citation)


def _add_table(doc: Document, rows: list) -> None:
    """Render a list-of-lists as a Word table with bold header row."""
    if not rows:
        return
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = str(cell_text)
            if row_idx == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()  # spacing after table


def export_sections_to_docx(
    sections: List[DraftSection],
    output_path: Path,
    locale: str = "en",
    client_branding: dict | None = None,
) -> Path:
    """
    Write sections to a DOCX file at output_path.

    Returns output_path.
    """
    output_path.parent.mkdir(parents=True, exist_ok=True)

    branding = client_branding or {}
    company_name: str = branding.get("company_name") or "Report Draft"
    primary_color: RGBColor | None = None
    if "primary_color" in branding:
        primary_color = _parse_rgb(branding["primary_color"])

    doc = Document()

    _add_title_page(doc, company_name, locale)

    for section in sections:
        h1 = doc.add_heading(section.heading, level=1)
        if primary_color:
            _apply_heading_color(h1, primary_color)

        for block in section.blocks:
            btype = block.block_type
            content = block.content

            if btype == "paragraph":
                doc.add_paragraph(str(content) if not isinstance(content, str) else content)

            elif btype == "heading":
                h2 = doc.add_heading(str(content) if not isinstance(content, str) else content, level=2)
                if primary_color:
                    _apply_heading_color(h2, primary_color)

            elif btype == "bullet_list":
                items = content if isinstance(content, list) else [content]
                for item in items:
                    _add_bullet(doc, item)

            elif btype == "table":
                rows = content if isinstance(content, list) else [[content]]
                _add_table(doc, rows)

            else:
                # Fallback: stringify and add as paragraph
                doc.add_paragraph(str(content))

        if section.facts:
            sub = doc.add_heading("Key Facts", level=2)
            if primary_color:
                _apply_heading_color(sub, primary_color)
            for fact in section.facts:
                _add_bullet(doc, fact)

        if section.citations:
            sub = doc.add_heading("References", level=2)
            if primary_color:
                _apply_heading_color(sub, primary_color)
            for citation in section.citations:
                _add_numbered(doc, _citation_text(citation))

    doc.save(str(output_path))
    return output_path
