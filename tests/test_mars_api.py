"""
FastAPI integration tests for MARS API endpoints.

Uses TestClient against the real app with an isolated SQLite database.
Migration runs during app startup (triggered by TestClient lifespan).
"""

from __future__ import annotations

import json
from pathlib import Path

import pytest

pytest.importorskip("fastapi")

from fastapi.testclient import TestClient  # noqa: E402
from freezegun import freeze_time  # noqa: E402

import db.connection as db_connection_module  # noqa: E402


@pytest.fixture()
def test_db_path(tmp_path: Path) -> Path:
    return tmp_path / "test_mars_api.db"


@pytest.fixture()
def client(test_db_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    """
    Create a TestClient with DB_PATH redirected to a temp database.
    Monkeypatch MUST happen before TestClient is created (startup runs migrations).
    """
    monkeypatch.setattr(db_connection_module, "DB_PATH", test_db_path)
    # Re-import after patch so startup sees the new path
    from api.main import app
    with TestClient(app, raise_server_exceptions=True) as c:
        yield c


# ── Locales ────────────────────────────────────────────────────────────────────

class TestLocalesRouter:
    def test_list_locales_returns_24(self, client: TestClient) -> None:
        resp = client.get("/api/locales")
        assert resp.status_code == 200
        data = resp.json()
        assert len(data) == 24

    def test_locale_has_required_fields(self, client: TestClient) -> None:
        resp = client.get("/api/locales")
        first = resp.json()[0]
        assert "code" in first
        assert "name" in first
        assert "native_name" in first
        assert "is_active" in first

    def test_get_client_locale_settings_default(self, client: TestClient) -> None:
        # gerold-foods is a real client in the registry
        resp = client.get("/api/clients/gerold-foods/locale-settings")
        assert resp.status_code == 200
        data = resp.json()
        assert data["client_id"] == "gerold-foods"
        assert data["fallback_locale"] == "en"

    def test_put_client_locale_settings(self, client: TestClient) -> None:
        resp = client.put(
            "/api/clients/gerold-foods/locale-settings",
            json={"primary_locale": "de", "enabled_locales": ["en", "de", "fr"]},
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["primary_locale"] == "de"
        assert "de" in data["enabled_locales"]

    def test_locale_settings_persists(self, client: TestClient) -> None:
        client.put(
            "/api/clients/gerold-foods/locale-settings",
            json={"primary_locale": "fr", "enabled_locales": ["en", "fr"]},
        )
        resp = client.get("/api/clients/gerold-foods/locale-settings")
        assert resp.json()["primary_locale"] == "fr"

    def test_locale_settings_unknown_client_404(self, client: TestClient) -> None:
        resp = client.get("/api/clients/nonexistent-xyz/locale-settings")
        assert resp.status_code == 404


# ── Keywords ───────────────────────────────────────────────────────────────────

class TestKeywordsRouter:
    def test_list_keywords_empty(self, client: TestClient) -> None:
        resp = client.get("/api/clients/gerold-foods/keywords")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_create_keyword(self, client: TestClient) -> None:
        resp = client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "packaging waste", "locale": "en", "weight": 2.0},
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["phrase"] == "packaging waste"
        assert data["weight"] == 2.0
        assert data["locale"] == "en"

    def test_create_keyword_duplicate_409(self, client: TestClient) -> None:
        client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "ghg reporting", "locale": "en"},
        )
        resp = client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "ghg reporting", "locale": "en"},
        )
        assert resp.status_code == 409

    def test_get_keyword(self, client: TestClient) -> None:
        created = client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "scope 3", "locale": "en"},
        ).json()
        resp = client.get(f"/api/clients/gerold-foods/keywords/{created['id']}")
        assert resp.status_code == 200
        assert resp.json()["phrase"] == "scope 3"

    def test_update_keyword(self, client: TestClient) -> None:
        created = client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "water use", "locale": "en", "weight": 1.0},
        ).json()
        resp = client.put(
            f"/api/clients/gerold-foods/keywords/{created['id']}",
            json={"phrase": "water usage", "locale": "en", "weight": 3.0},
        )
        assert resp.status_code == 200
        assert resp.json()["phrase"] == "water usage"
        assert resp.json()["weight"] == 3.0

    def test_delete_keyword(self, client: TestClient) -> None:
        created = client.post(
            "/api/clients/gerold-foods/keywords",
            json={"phrase": "to delete", "locale": "en"},
        ).json()
        resp = client.delete(f"/api/clients/gerold-foods/keywords/{created['id']}")
        assert resp.status_code == 200
        # Should no longer appear in list
        listed = client.get("/api/clients/gerold-foods/keywords").json()
        assert all(k["id"] != created["id"] for k in listed)

    def test_keyword_preview(self, client: TestClient) -> None:
        resp = client.post(
            "/api/clients/gerold-foods/keywords/preview",
            json={
                "phrases": ["packaging", "water"],
                "locale": "en",
                "sample_text": "New packaging regulation affects water treatment facilities.",
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["match_count"] == 2
        phrases_found = {m["phrase"] for m in data["matches"]}
        assert "packaging" in phrases_found
        assert "water" in phrases_found

    def test_keyword_unknown_client_404(self, client: TestClient) -> None:
        resp = client.get("/api/clients/no-such-client/keywords")
        assert resp.status_code == 404


# ── Source proposals ───────────────────────────────────────────────────────────

class TestSourcesRouter:
    def test_list_proposals_empty(self, client: TestClient) -> None:
        resp = client.get("/api/clients/gerold-foods/source-proposals")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_suggest_sources_returns_list(self, client: TestClient) -> None:
        resp = client.post("/api/clients/gerold-foods/source-proposals/suggest")
        assert resp.status_code == 201
        assert isinstance(resp.json(), list)

    def test_suggest_sources_reuses_existing_pending_proposal(self, client: TestClient) -> None:
        first = client.post("/api/clients/gerold-foods/source-proposals/suggest")
        assert first.status_code == 201
        initial = first.json()
        assert len(initial) == 1

        second = client.post("/api/clients/gerold-foods/source-proposals/suggest")
        assert second.status_code == 201
        repeated = second.json()
        assert len(repeated) == 1
        assert repeated[0]["id"] == initial[0]["id"]

        listed = client.get("/api/clients/gerold-foods/source-proposals")
        assert listed.status_code == 200
        assert len(listed.json()) == 1

    def test_approve_proposal(self, client: TestClient) -> None:
        # Manually insert a proposal via DB then approve via API
        import db.connection as dbc
        with dbc.get_db() as conn:
            cur = conn.execute(
                """INSERT INTO source_proposals (client_id, url, title, proposed_by)
                   VALUES (?, ?, ?, ?)""",
                ("gerold-foods", "https://eur-lex.europa.eu/", "EUR-Lex", "source-scout"),
            )
            proposal_id = cur.lastrowid

        resp = client.post(
            f"/api/clients/gerold-foods/source-proposals/{proposal_id}/action",
            json={"action": "approve", "reviewer": "alice"},
        )
        assert resp.status_code == 200
        assert resp.json()["status"] == "approved"
        assert resp.json()["reviewed_by"] == "alice"

    def test_reject_proposal(self, client: TestClient) -> None:
        import db.connection as dbc
        with dbc.get_db() as conn:
            cur = conn.execute(
                """INSERT INTO source_proposals (client_id, url, proposed_by)
                   VALUES (?, ?, ?)""",
                ("gerold-foods", "https://example.com", "source-scout"),
            )
            proposal_id = cur.lastrowid

        resp = client.post(
            f"/api/clients/gerold-foods/source-proposals/{proposal_id}/action",
            json={"action": "reject", "reviewer": "bob", "note": "Duplicate source"},
        )
        assert resp.status_code == 200
        assert resp.json()["status"] == "rejected"

    def test_filter_by_status(self, client: TestClient) -> None:
        import db.connection as dbc
        with dbc.get_db() as conn:
            conn.execute(
                """INSERT INTO source_proposals (client_id, url, status, proposed_by)
                   VALUES (?, ?, ?, ?)""",
                ("gerold-foods", "https://a.com", "approved", "scout"),
            )
            conn.execute(
                """INSERT INTO source_proposals (client_id, url, status, proposed_by)
                   VALUES (?, ?, ?, ?)""",
                ("gerold-foods", "https://b.com", "pending", "scout"),
            )

        resp = client.get("/api/clients/gerold-foods/source-proposals?status=approved")
        assert resp.status_code == 200
        results = resp.json()
        assert all(r["status"] == "approved" for r in results)


# ── Templates ──────────────────────────────────────────────────────────────────

class TestTemplatesRouter:
    def test_list_templates_empty(self, client: TestClient) -> None:
        resp = client.get("/api/templates")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_create_template(self, client: TestClient) -> None:
        resp = client.post(
            "/api/templates",
            json={"name": "CSRD Q-Report", "base_locale": "en"},
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["name"] == "CSRD Q-Report"
        assert data["current_version"] == 0

    def test_get_template(self, client: TestClient) -> None:
        created = client.post(
            "/api/templates",
            json={"name": "T1", "base_locale": "de"},
        ).json()
        resp = client.get(f"/api/templates/{created['id']}")
        assert resp.status_code == 200
        assert resp.json()["base_locale"] == "de"

    def test_update_template(self, client: TestClient) -> None:
        created = client.post(
            "/api/templates",
            json={"name": "Old Name", "base_locale": "en"},
        ).json()
        resp = client.put(
            f"/api/templates/{created['id']}",
            json={"name": "New Name", "base_locale": "en"},
        )
        assert resp.status_code == 200
        assert resp.json()["name"] == "New Name"

    def test_delete_template(self, client: TestClient) -> None:
        created = client.post(
            "/api/templates",
            json={"name": "Delete Me", "base_locale": "en"},
        ).json()
        resp = client.delete(f"/api/templates/{created['id']}")
        assert resp.status_code == 200
        assert client.get(f"/api/templates/{created['id']}").status_code == 404

    def test_create_version(self, client: TestClient) -> None:
        tmpl = client.post(
            "/api/templates",
            json={"name": "Versioned", "base_locale": "en"},
        ).json()
        resp = client.post(
            f"/api/templates/{tmpl['id']}/versions",
            json={
                "sections_json": [{"section_id": "s1", "heading": "Intro"}],
                "changelog_note": "v2",
            },
        )
        assert resp.status_code == 201
        assert resp.json()["version_number"] == 1
        assert resp.json()["changelog_note"] == "v2"

    def test_list_versions(self, client: TestClient) -> None:
        tmpl = client.post(
            "/api/templates",
            json={"name": "Multi", "base_locale": "en"},
        ).json()
        client.post(f"/api/templates/{tmpl['id']}/versions", json={})
        client.post(f"/api/templates/{tmpl['id']}/versions", json={})
        resp = client.get(f"/api/templates/{tmpl['id']}/versions")
        assert resp.status_code == 200
        assert len(resp.json()) == 2

    def test_clone_template(self, client: TestClient) -> None:
        src = client.post(
            "/api/templates",
            json={"name": "Source Template", "base_locale": "en"},
        ).json()
        resp = client.post(
            f"/api/templates/{src['id']}/clone",
            json={"name": "Cloned Template"},
        )
        assert resp.status_code == 201
        cloned = resp.json()
        assert cloned["name"] == "Cloned Template"
        assert cloned["id"] != src["id"]
        assert cloned["base_locale"] == src["base_locale"]

    def test_update_template_sections_creates_new_version(self, client: TestClient) -> None:
        tmpl = client.post(
            "/api/templates",
            json={"name": "Section Save", "base_locale": "en"},
        ).json()
        resp = client.put(
            f"/api/templates/{tmpl['id']}/sections",
            json={
                "sections": [
                    {
                        "section_id": "executive_summary",
                        "heading": "Executive Summary",
                        "order": 0,
                        "prompt_template": "Summarize changes.",
                        "chart_types": ["topic_distribution"],
                        "max_tokens": 900,
                        "required": True,
                    }
                ],
                "changelog": "Initial section setup",
            },
        )
        assert resp.status_code == 200
        version = resp.json()
        assert version["version_number"] == 1
        assert version["sections_json"][0]["heading"] == "Executive Summary"

    def test_update_template_theme_creates_new_version(self, client: TestClient) -> None:
        tmpl = client.post(
            "/api/templates",
            json={"name": "Theme Save", "base_locale": "en"},
        ).json()
        client.post(
            f"/api/templates/{tmpl['id']}/versions",
            json={
                "sections_json": [{"section_id": "s1", "heading": "Intro"}],
                "theme_tokens": {"brand_color": "#112233"},
            },
        )
        resp = client.put(
            f"/api/templates/{tmpl['id']}/theme",
            json={"tokens": {"brand_color": "#224466", "surface_color": "#f8f8f8"}},
        )
        assert resp.status_code == 200
        assert resp.json()["current_version"] == 2

    def test_template_not_found_404(self, client: TestClient) -> None:
        resp = client.get("/api/templates/9999")
        assert resp.status_code == 404


# ── Drafts ─────────────────────────────────────────────────────────────────────

class TestDraftsRouter:
    def _create_draft(self, client: TestClient, title: str = "Test Draft") -> dict:
        return client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": title, "primary_locale": "en"},
        ).json()

    def test_list_drafts_empty(self, client: TestClient) -> None:
        resp = client.get("/api/clients/gerold-foods/drafts")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_create_draft(self, client: TestClient) -> None:
        resp = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Feb 2026 Report", "primary_locale": "en", "period": "2026-02"},
        )
        assert resp.status_code == 201
        data = resp.json()
        assert data["title"] == "Feb 2026 Report"
        assert data["status"] == "composing"
        assert data["primary_locale"] == "en"

    def test_get_draft_detail(self, client: TestClient) -> None:
        created = self._create_draft(client)
        resp = client.get(f"/api/clients/gerold-foods/drafts/{created['id']}")
        assert resp.status_code == 200
        assert resp.json()["id"] == created["id"]

    def test_compose_creates_revision(self, client: TestClient) -> None:
        created = self._create_draft(client)
        resp = client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        assert resp.status_code == 201
        rev = resp.json()
        assert rev["revision_number"] == 1
        assert "sections" in rev
        assert len(rev["sections"]) > 0

    def test_compose_increments_revision(self, client: TestClient) -> None:
        created = self._create_draft(client)
        client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        resp = client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        assert resp.status_code == 201
        assert resp.json()["revision_number"] == 2

    def test_list_revisions(self, client: TestClient) -> None:
        created = self._create_draft(client)
        client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        resp = client.get(f"/api/clients/gerold-foods/drafts/{created['id']}/revisions")
        assert resp.status_code == 200
        assert len(resp.json()) == 2

    def test_transition_composing_to_review(self, client: TestClient) -> None:
        created = self._create_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{created['id']}/transition",
            params={"target_status": "review"},
        )
        assert resp.status_code == 200
        assert resp.json()["status"] == "review"

    def test_invalid_transition_422(self, client: TestClient) -> None:
        created = self._create_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{created['id']}/transition",
            params={"target_status": "approved"},
        )
        assert resp.status_code == 422

    def test_archive_draft(self, client: TestClient) -> None:
        created = self._create_draft(client)
        resp = client.delete(f"/api/clients/gerold-foods/drafts/{created['id']}")
        assert resp.status_code == 200
        detail = client.get(f"/api/clients/gerold-foods/drafts/{created['id']}").json()
        assert detail["status"] == "archived"

    def test_draft_not_found_404(self, client: TestClient) -> None:
        resp = client.get("/api/clients/gerold-foods/drafts/9999")
        assert resp.status_code == 404

    def test_translate_creates_revision(self, client: TestClient) -> None:
        created = self._create_draft(client)
        client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{created['id']}/translate",
            params={"target_locale": "de"},
        )
        assert resp.status_code == 201
        rev = resp.json()
        assert rev["revision_number"] == 2
        for section in rev["sections"]:
            assert section["locale"] == "de"
            assert section["translation_status"] in {"translated", "low_confidence"}

    def test_update_section_creates_new_revision(self, client: TestClient) -> None:
        created = self._create_draft(client)
        rev = client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose").json()
        section_id = rev["sections"][0]["section_id"]
        resp = client.put(
            f"/api/clients/gerold-foods/drafts/{created['id']}/sections/{section_id}",
            json={
                "blocks": [{"type": "paragraph", "content": "Updated by editor."}],
                "revision_note": "Manual edit",
            },
        )
        assert resp.status_code == 200
        assert resp.json()["blocks"][0]["content"] == "Updated by editor."
        detail = client.get(f"/api/clients/gerold-foods/drafts/{created['id']}").json()
        assert detail["current_revision"]["revision_number"] == 2

    def test_translate_invalid_locale_422(self, client: TestClient) -> None:
        created = self._create_draft(client)
        client.post(f"/api/clients/gerold-foods/drafts/{created['id']}/compose")
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{created['id']}/translate",
            params={"target_locale": "xx"},
        )
        assert resp.status_code == 422

    def test_compose_downgrades_approved_draft_back_to_composing(self, client: TestClient) -> None:
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Needs Reapproval", "primary_locale": "en"},
        ).json()
        revision = client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose").json()
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "review"},
        )
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "approval"},
        )
        for section in revision["sections"]:
            client.post(
                f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
                json={
                    "section_id": section["section_id"],
                    "status": "approved",
                    "reviewer": "alice",
                },
            )
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        assert detail["status"] == "approved"

        resp = client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose")
        assert resp.status_code == 201
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        assert detail["status"] == "composing"
        assert client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}/approvals").json() == []


# ── Chat ───────────────────────────────────────────────────────────────────────

class TestChatRouter:
    def _setup_draft(self, client: TestClient) -> dict:
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Chat Test Draft", "primary_locale": "en"},
        ).json()
        client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose")
        return draft

    def test_list_chat_empty(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        resp = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}/chat")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_send_chat_message(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "Please expand section 1.", "section_id": "executive_summary"},
        )
        assert resp.status_code == 201
        msg = resp.json()
        assert msg["role"] == "assistant"
        assert msg["section_id"] == "executive_summary"

    def test_send_chat_message_creates_revision_and_assistant_reply(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        before = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "Please expand section 1.", "section_id": "executive_summary"},
        )
        assert resp.status_code == 201
        after = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        assert after["current_revision_id"] != before["current_revision_id"]
        messages = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}/chat").json()
        assert len(messages) == 2
        assert [message["role"] for message in messages] == ["user", "assistant"]

    def test_filter_chat_by_section(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "Expand intro.", "section_id": "executive_summary"},
        )
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/chat",
            json={"role": "user", "content": "Draft-level question."},
        )
        resp = client.get(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/chat",
            params={"section_id": "executive_summary"},
        )
        messages = resp.json()
        assert len(messages) == 2
        assert all(message["section_id"] == "executive_summary" for message in messages)


# ── Approvals ──────────────────────────────────────────────────────────────────

class TestApprovalsRouter:
    def _setup_draft(self, client: TestClient) -> dict:
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Approval Draft", "primary_locale": "en"},
        ).json()
        client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose")
        return draft

    def test_list_approvals_empty(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        resp = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}/approvals")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_approve_section(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
            json={
                "section_id": "executive_summary",
                "status": "approved",
                "reviewer": "alice",
                "comment": "Looks good",
            },
        )
        assert resp.status_code == 200
        data = resp.json()
        assert data["status"] == "approved"
        assert data["reviewer"] == "alice"
        assert data["section_id"] == "executive_summary"

    def test_reject_section(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
            json={
                "section_id": "regulatory_changes",
                "status": "rejected",
                "reviewer": "bob",
            },
        )
        assert resp.status_code == 200
        assert resp.json()["status"] == "rejected"

    def test_upsert_approval_updates_existing(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        real_section_id = detail["current_revision"]["sections"][0]["section_id"]
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
            json={"section_id": real_section_id, "status": "needs_revision", "reviewer": "alice"},
        )
        # Update to approved
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
            json={"section_id": real_section_id, "status": "approved", "reviewer": "alice"},
        )
        assert resp.json()["status"] == "approved"
        # Still only one record
        all_approvals = client.get(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/approvals"
        ).json()
        matching = [a for a in all_approvals if a["section_id"] == real_section_id]
        assert len(matching) == 1

    def test_all_sections_approved_rolls_draft_to_approved(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "approval"},
        )
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        for section in detail["current_revision"]["sections"]:
            client.post(
                f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
                json={
                    "section_id": section["section_id"],
                    "status": "approved",
                    "reviewer": "alice",
                },
            )
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        assert detail["status"] == "approved"

    def test_transition_to_approved_requires_section_approvals(self, client: TestClient) -> None:
        draft = self._setup_draft(client)
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "review"},
        )
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "approval"},
        )
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "approved"},
        )
        assert resp.status_code == 422


# ── Exports ────────────────────────────────────────────────────────────────────

class TestExportsRouter:
    def _approved_draft(self, client: TestClient) -> dict:
        """Create a draft, compose it, and approve every section."""
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Export Draft", "primary_locale": "en"},
        ).json()
        revision = client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose").json()
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "review"},
        )
        client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/transition",
            params={"target_status": "approval"},
        )
        for section in revision["sections"]:
            client.post(
                f"/api/clients/gerold-foods/drafts/{draft['id']}/approve",
                json={
                    "section_id": section["section_id"],
                    "status": "approved",
                    "reviewer": "alice",
                },
            )
        detail = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}").json()
        assert detail["status"] == "approved"
        return draft

    def test_list_exports_empty(self, client: TestClient) -> None:
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "D", "primary_locale": "en"},
        ).json()
        resp = client.get(f"/api/clients/gerold-foods/drafts/{draft['id']}/exports")
        assert resp.status_code == 200
        assert resp.json() == []

    def test_request_docx_export_without_approval(self, client: TestClient) -> None:
        """DOCX export completes immediately for unapproved drafts."""
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "DOCX Draft", "primary_locale": "en"},
        ).json()
        client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose")
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports",
            json={"format": "docx", "locale": "en"},
        )
        assert resp.status_code == 201
        assert resp.json()["format"] == "docx"
        assert resp.json()["status"] == "completed"

    def test_request_pdf_export_requires_approval(self, client: TestClient) -> None:
        """PDF export on a composing draft returns 403."""
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "PDF Blocked", "primary_locale": "en"},
        ).json()
        client.post(f"/api/clients/gerold-foods/drafts/{draft['id']}/compose")
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 403

    def test_request_pdf_export_approved_draft(self, client: TestClient) -> None:
        """PDF export completes immediately for approved draft."""
        draft = self._approved_draft(client)
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports",
            json={"format": "pdf", "locale": "en"},
        )
        assert resp.status_code == 201
        job = resp.json()
        assert job["format"] == "pdf"
        assert job["status"] == "completed"

    def test_get_export_job(self, client: TestClient) -> None:
        draft = self._approved_draft(client)
        job = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports",
            json={"format": "pdf"},
        ).json()
        resp = client.get(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports/{job['id']}"
        )
        assert resp.status_code == 200
        assert resp.json()["id"] == job["id"]

    def test_download_not_ready_404(self, client: TestClient) -> None:
        """PDF export becomes downloadable immediately after synchronous processing."""
        draft = self._approved_draft(client)
        job = client.post(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports",
            json={"format": "pdf"},
        ).json()
        status = client.get(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports/{job['id']}"
        ).json()
        assert status["status"] == "completed"
        resp = client.get(
            f"/api/clients/gerold-foods/drafts/{draft['id']}/exports/{job['id']}/download"
        )
        assert resp.status_code == 200

    def test_export_rejects_revision_from_another_draft(self, client: TestClient) -> None:
        approved = self._approved_draft(client)
        other = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Other Draft", "primary_locale": "en"},
        ).json()
        other_rev = client.post(f"/api/clients/gerold-foods/drafts/{other['id']}/compose").json()
        resp = client.post(
            f"/api/clients/gerold-foods/drafts/{approved['id']}/exports",
            json={"format": "pdf", "revision_id": other_rev["id"]},
        )
        assert resp.status_code == 422


# ── DOCX import ────────────────────────────────────────────────────────────────

class TestDocxImport:
    def test_import_docx_creates_revision(self, client: TestClient, tmp_path: Path) -> None:
        """Upload a minimal DOCX and verify a new revision is created."""
        from docx import Document
        draft = client.post(
            "/api/clients/gerold-foods/drafts",
            json={"title": "Import Draft", "primary_locale": "en"},
        ).json()

        # Create a minimal DOCX
        doc = Document()
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph("This section describes key findings.")
        doc.add_heading("Actions Required", level=1)
        doc.add_paragraph("Complete packaging audit.")
        docx_path = tmp_path / "test.docx"
        doc.save(str(docx_path))

        with open(docx_path, "rb") as f:
            resp = client.post(
                f"/api/clients/gerold-foods/drafts/{draft['id']}/exports/import-docx",
                files={"file": ("test.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
        assert resp.status_code == 201
        data = resp.json()
        assert data["draft_id"] == draft["id"]
        assert data["sections_imported"] == 2
        assert data["revision_id"] > 0
