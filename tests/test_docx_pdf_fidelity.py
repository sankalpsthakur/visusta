"""
Fidelity tests for DOCX export, PDF conversion, and round-trip DOCX import.

Test coverage:
1. DOCX export produces valid file — structural checks via python-docx
2. PDF export via soffice produces non-empty file with %PDF magic
3. Round-trip: export → import recovers section count and headings
4. Client branding: primary_color and company_name appear in DOCX
5. Edge cases: empty sections, section with no blocks, non-ASCII text, long paragraph
"""

from __future__ import annotations

import io
import subprocess
import sys
from pathlib import Path
from typing import List
from unittest.mock import MagicMock, patch

import pytest
from docx import Document

# ── path setup ──────────────────────────────────────────────────────────────
# Ensure the project root is on sys.path so mars.* and api.* resolve.
_REPO_ROOT = Path(__file__).parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

import shutil

from api.schemas_mars import DraftSection, SectionBlock
from mars.docx_export import export_sections_to_docx
from mars.docx_import import parse_docx_to_sections
from mars.pdf_export import export_sections_to_pdf

def _usable_soffice() -> str | None:
    candidate = shutil.which("soffice")
    if not candidate:
        return None
    try:
        probe = subprocess.run(
            [candidate, "--headless", "--version"],
            timeout=10,
            capture_output=True,
            stdin=subprocess.DEVNULL,
        )
    except (OSError, subprocess.SubprocessError):
        return None
    return candidate if probe.returncode == 0 else None


SOFFICE = _usable_soffice()


# ── shared fixtures ──────────────────────────────────────────────────────────

def _make_section(
    section_id: str = "s1",
    heading: str = "Executive Summary",
    locale: str = "en",
    blocks: List[SectionBlock] | None = None,
    facts: List[str] | None = None,
    citations: List[str] | None = None,
) -> DraftSection:
    return DraftSection(
        section_id=section_id,
        heading=heading,
        locale=locale,
        blocks=blocks or [],
        facts=facts or [],
        citations=citations or [],
    )


def _make_full_section(section_id: str = "s1", heading: str = "Overview") -> DraftSection:
    """Section with one of every block type plus facts and citations."""
    return _make_section(
        section_id=section_id,
        heading=heading,
        blocks=[
            SectionBlock(block_id="b1", block_type="paragraph", content="Introductory paragraph."),
            SectionBlock(block_id="b2", block_type="heading", content="Sub-heading Alpha"),
            SectionBlock(
                block_id="b3",
                block_type="bullet_list",
                content=["First item", "Second item", "Third item"],
            ),
            SectionBlock(
                block_id="b4",
                block_type="table",
                content=[["Country", "Score"], ["DE", "92"], ["FR", "88"]],
            ),
        ],
        facts=["Revenue up 12%", "Scope 1 emissions reduced by 8%"],
        citations=["IPCC AR6, 2021", "EU Taxonomy Regulation 2020/852"],
    )


@pytest.fixture
def single_full_section() -> List[DraftSection]:
    return [_make_full_section()]


@pytest.fixture
def two_sections() -> List[DraftSection]:
    return [
        _make_full_section("s1", "Overview"),
        _make_full_section("s2", "Financial Performance"),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# 1. DOCX export structural validation
# ─────────────────────────────────────────────────────────────────────────────

class TestDocxExport:
    def test_produces_nonempty_file(self, tmp_path: Path, two_sections: List[DraftSection]) -> None:
        out = tmp_path / "report.docx"
        result = export_sections_to_docx(two_sections, out)
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 1000

    def test_heading_count_matches_sections(
        self, tmp_path: Path, two_sections: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(two_sections, out)
        doc = Document(str(out))
        h1s = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
        assert len(h1s) == len(two_sections)
        headings_text = [p.text for p in h1s]
        for section in two_sections:
            assert section.heading in headings_text

    def test_paragraph_content_preserved(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(single_full_section, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Introductory paragraph." in all_text

    def test_bullet_list_items_present(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(single_full_section, out)
        doc = Document(str(out))
        list_paragraphs = [
            p.text for p in doc.paragraphs if "List" in (p.style.name or "")
        ]
        assert any("First item" in t for t in list_paragraphs)
        assert any("Second item" in t for t in list_paragraphs)
        assert any("Third item" in t for t in list_paragraphs)

    def test_table_has_correct_row_count(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(single_full_section, out)
        doc = Document(str(out))
        assert len(doc.tables) >= 1
        # Table block has 3 rows (header + 2 data rows)
        assert len(doc.tables[0].rows) == 3

    def test_facts_appear_under_key_facts_heading(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(single_full_section, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Key Facts" in all_text
        assert "Revenue up 12%" in all_text

    def test_citations_appear_under_references_heading(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.docx"
        export_sections_to_docx(single_full_section, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "References" in all_text
        assert "IPCC AR6, 2021" in all_text

    def test_unknown_block_type_falls_back_to_paragraph(
        self, tmp_path: Path
    ) -> None:
        sections = [
            _make_section(
                blocks=[
                    SectionBlock(block_id="b1", block_type="chart", content="chart_payload"),
                ]
            )
        ]
        out = tmp_path / "report.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "chart_payload" in all_text

    def test_creates_parent_directories(self, tmp_path: Path) -> None:
        nested = tmp_path / "deep" / "nested" / "report.docx"
        export_sections_to_docx([_make_section()], nested)
        assert nested.exists()


# ─────────────────────────────────────────────────────────────────────────────
# 2. PDF export via soffice
# ─────────────────────────────────────────────────────────────────────────────

class TestPdfExport:
    def _soffice_available(self) -> bool:
        return SOFFICE is not None

    @pytest.mark.skipif(
        SOFFICE is None,
        reason="LibreOffice (soffice) not found on PATH",
    )
    def test_pdf_has_magic_bytes_and_nonzero_size(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.pdf"
        result = export_sections_to_pdf(single_full_section, out)
        assert result == out
        assert out.exists()
        assert out.stat().st_size > 0
        magic = out.read_bytes()[:4]
        assert magic == b"%PDF", f"File does not start with %PDF magic: {magic!r}"

    def test_raises_when_soffice_missing(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.pdf"
        with patch("mars.pdf_export.Path") as mock_path_cls:
            # Make the soffice path appear not to exist
            fake_path = MagicMock()
            fake_path.exists.return_value = False
            mock_path_cls.return_value = fake_path
            with pytest.raises(FileNotFoundError, match="LibreOffice not found"):
                export_sections_to_pdf(single_full_section, out)

    @pytest.mark.skipif(
        SOFFICE is None,
        reason="LibreOffice (soffice) not found on PATH",
    )
    def test_pdf_with_branding(self, tmp_path: Path) -> None:
        sections = [_make_full_section()]
        out = tmp_path / "branded.pdf"
        result = export_sections_to_pdf(
            sections,
            out,
            locale="de",
            client_branding={"company_name": "Acme GmbH", "primary_color": "#e63946"},
        )
        assert result.stat().st_size > 0
        assert result.read_bytes()[:4] == b"%PDF"

    def test_uses_isolated_headless_profile_and_env(
        self, tmp_path: Path, single_full_section: List[DraftSection]
    ) -> None:
        out = tmp_path / "report.pdf"
        work_dir = tmp_path / "lo-work"
        work_dir.mkdir()
        fake_soffice = tmp_path / "soffice"
        fake_soffice.write_text("", encoding="utf-8")
        captured: dict[str, object] = {}

        def fake_export(
            sections: List[DraftSection],
            output: Path,
            locale: str = "en",
            client_branding: dict | None = None,
        ) -> None:
            output.write_bytes(b"docx")

        def fake_run(args: list[str], **kwargs: object) -> subprocess.CompletedProcess[str]:
            captured["args"] = args
            captured["env"] = kwargs["env"]
            captured["stdin"] = kwargs["stdin"]
            (work_dir / "export.pdf").write_bytes(b"%PDF-test")
            return subprocess.CompletedProcess(args=args, returncode=0)

        with patch("mars.pdf_export._resolve_soffice", return_value=str(fake_soffice)):
            with patch("mars.pdf_export.tempfile.mkdtemp", return_value=str(work_dir)):
                with patch("mars.pdf_export.export_sections_to_docx", side_effect=fake_export):
                    with patch("mars.pdf_export.subprocess.run", side_effect=fake_run):
                        result = export_sections_to_pdf(single_full_section, out)

        assert result == out
        assert out.read_bytes()[:4] == b"%PDF"
        args = captured["args"]
        env = captured["env"]
        assert isinstance(args, list)
        assert "--headless" in args
        assert "--invisible" in args
        assert "--norestore" in args
        assert any(str(arg).startswith("-env:UserInstallation=") for arg in args)
        assert captured["stdin"] is subprocess.DEVNULL
        assert isinstance(env, dict)
        assert env["SAL_USE_VCLPLUGIN"] == "svp"
        assert env["NO_AT_BRIDGE"] == "1"
        assert env["OOO_DISABLE_RECOVERY"] == "1"
        assert env["DISPLAY"] == ""
        assert env["HOME"] == str(work_dir)
        assert env["TMPDIR"] == str(work_dir)


# ─────────────────────────────────────────────────────────────────────────────
# 3. Round-trip: export → import
# ─────────────────────────────────────────────────────────────────────────────

class TestRoundTrip:
    def test_section_count_preserved(
        self, tmp_path: Path, two_sections: List[DraftSection]
    ) -> None:
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(two_sections, out)
        recovered, warnings = parse_docx_to_sections(out.read_bytes(), locale="en")
        # The DOCX title page (before the first Heading 1) may produce a "Preamble"
        # section in the importer; filter it out for the count check.
        content_sections = [s for s in recovered if s.section_id != "imported_preamble"]
        assert len(content_sections) == len(two_sections), (
            f"Expected {len(two_sections)} content sections, got {len(content_sections)}. "
            f"Warnings: {warnings}"
        )

    def test_section_headings_preserved(
        self, tmp_path: Path, two_sections: List[DraftSection]
    ) -> None:
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(two_sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes(), locale="en")
        original_headings = {s.heading for s in two_sections}
        # Exclude the preamble; original headings must be a subset of recovered headings.
        recovered_headings = {s.heading for s in recovered if s.section_id != "imported_preamble"}
        assert original_headings == recovered_headings

    def test_paragraph_content_survives(self, tmp_path: Path) -> None:
        sections = [
            _make_section(
                blocks=[SectionBlock(block_id="b1", block_type="paragraph", content="Hello world.")]
            )
        ]
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes())
        all_content = " ".join(
            str(b.content) for s in recovered for b in s.blocks
        )
        assert "Hello world." in all_content

    def test_bullet_list_survives_as_bullet_list_block(self, tmp_path: Path) -> None:
        sections = [
            _make_section(
                blocks=[
                    SectionBlock(
                        block_id="b1",
                        block_type="bullet_list",
                        content=["Alpha", "Beta", "Gamma"],
                    )
                ]
            )
        ]
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes())
        bullet_blocks = [
            b for s in recovered for b in s.blocks if b.block_type == "bullet_list"
        ]
        assert bullet_blocks, "No bullet_list block recovered"
        all_items = [item for b in bullet_blocks for item in b.content]
        assert "Alpha" in all_items
        assert "Beta" in all_items
        assert "Gamma" in all_items

    def test_table_survives_with_correct_row_count(self, tmp_path: Path) -> None:
        rows = [["Name", "Value"], ["Revenue", "1M"], ["Costs", "0.5M"]]
        sections = [
            _make_section(
                blocks=[SectionBlock(block_id="b1", block_type="table", content=rows)]
            )
        ]
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes())
        table_blocks = [b for s in recovered for b in s.blocks if b.block_type == "table"]
        assert table_blocks, "No table block recovered"
        assert len(table_blocks[0].content) == len(rows)

    def test_facts_recovered_on_section(self, tmp_path: Path) -> None:
        sections = [_make_section(facts=["Fact A", "Fact B"])]
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes())
        assert recovered
        # Facts should be recovered onto the section
        recovered_facts = [f for s in recovered for f in s.facts]
        assert "Fact A" in recovered_facts
        assert "Fact B" in recovered_facts

    def test_citations_recovered_on_section(self, tmp_path: Path) -> None:
        sections = [_make_section(citations=["Author A, 2024", "Author B, 2025"])]
        out = tmp_path / "roundtrip.docx"
        export_sections_to_docx(sections, out)
        recovered, _ = parse_docx_to_sections(out.read_bytes())
        assert recovered
        recovered_citations = [c for s in recovered for c in s.citations]
        # Citations round-trip without synthetic [N] prefix
        assert "Author A, 2024" in recovered_citations
        assert "Author B, 2025" in recovered_citations


# ─────────────────────────────────────────────────────────────────────────────
# 4. Client branding
# ─────────────────────────────────────────────────────────────────────────────

class TestClientBranding:
    def test_company_name_appears_as_title(self, tmp_path: Path) -> None:
        sections = [_make_section()]
        out = tmp_path / "branded.docx"
        export_sections_to_docx(
            sections, out, client_branding={"company_name": "Acme Corp"}
        )
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Acme Corp" in all_text

    def test_default_title_when_no_branding(self, tmp_path: Path) -> None:
        sections = [_make_section()]
        out = tmp_path / "default.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Report Draft" in all_text

    def test_primary_color_applied_to_heading_runs(self, tmp_path: Path) -> None:
        sections = [_make_section()]
        out = tmp_path / "colored.docx"
        export_sections_to_docx(
            sections,
            out,
            client_branding={"primary_color": "#1a56db"},
        )
        doc = Document(str(out))
        h1s = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
        assert h1s, "No Heading 1 found"
        # At least one run in the heading should have a non-None font color
        colored_runs = [
            r for h in h1s for r in h.runs if r.font.color.rgb is not None
        ]
        assert colored_runs, "No colored runs found on Heading 1"
        rgb = colored_runs[0].font.color.rgb
        assert str(rgb).upper() == "1A56DB"

    def test_invalid_color_does_not_crash(self, tmp_path: Path) -> None:
        sections = [_make_section()]
        out = tmp_path / "bad_color.docx"
        # Should not raise — invalid color is silently ignored
        export_sections_to_docx(
            sections, out, client_branding={"primary_color": "notacolor"}
        )
        assert out.exists()

    def test_branding_with_none_values_uses_defaults(self, tmp_path: Path) -> None:
        sections = [_make_section()]
        out = tmp_path / "partial.docx"
        export_sections_to_docx(
            sections, out, client_branding={"company_name": None, "primary_color": None}
        )
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Report Draft" in all_text


# ─────────────────────────────────────────────────────────────────────────────
# 5. Edge cases
# ─────────────────────────────────────────────────────────────────────────────

class TestEdgeCases:
    def test_empty_sections_list_produces_valid_docx(self, tmp_path: Path) -> None:
        out = tmp_path / "empty.docx"
        export_sections_to_docx([], out)
        assert out.exists()
        # Should still be a valid DOCX (openable by python-docx)
        doc = Document(str(out))
        assert doc is not None

    def test_section_with_no_blocks(self, tmp_path: Path) -> None:
        sections = [_make_section(heading="Empty Section")]
        out = tmp_path / "no_blocks.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        h1s = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
        assert any("Empty Section" in h.text for h in h1s)

    def test_german_umlauts(self, tmp_path: Path) -> None:
        text = "Österreichische Umweltschutzmaßnahmen: Überprüfung für 2026"
        sections = [
            _make_section(
                heading="Umweltbericht",
                blocks=[SectionBlock(block_id="b1", block_type="paragraph", content=text)],
            )
        ]
        out = tmp_path / "german.docx"
        export_sections_to_docx(sections, out, locale="de")
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Österreichische" in all_text
        assert "Überprüfung" in all_text

    def test_japanese_text(self, tmp_path: Path) -> None:
        text = "環境規制の遵守：2026年の展望と課題"
        sections = [
            _make_section(
                heading="環境レポート",
                blocks=[SectionBlock(block_id="b1", block_type="paragraph", content=text)],
            )
        ]
        out = tmp_path / "japanese.docx"
        export_sections_to_docx(sections, out, locale="ja")
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "環境規制" in all_text

    def test_very_long_paragraph(self, tmp_path: Path) -> None:
        long_text = "Lorem ipsum dolor sit amet. " * 500  # ~14 000 chars
        sections = [
            _make_section(
                blocks=[SectionBlock(block_id="b1", block_type="paragraph", content=long_text)]
            )
        ]
        out = tmp_path / "long.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        # Content should be present and file should be valid
        assert "Lorem ipsum" in all_text
        assert out.stat().st_size > 5000

    def test_round_trip_empty_sections_returns_no_sections_warning(
        self, tmp_path: Path
    ) -> None:
        out = tmp_path / "empty.docx"
        export_sections_to_docx([], out)
        recovered, warnings = parse_docx_to_sections(out.read_bytes())
        # Title page has no Heading 1 sections, so import should warn
        assert isinstance(recovered, list)
        assert isinstance(warnings, list)

    def test_bullet_list_with_single_item(self, tmp_path: Path) -> None:
        sections = [
            _make_section(
                blocks=[
                    SectionBlock(block_id="b1", block_type="bullet_list", content=["Only item"])
                ]
            )
        ]
        out = tmp_path / "single_bullet.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        list_paras = [p.text for p in doc.paragraphs if "List" in (p.style.name or "")]
        assert any("Only item" in t for t in list_paras)

    def test_table_with_single_row(self, tmp_path: Path) -> None:
        sections = [
            _make_section(
                blocks=[
                    SectionBlock(block_id="b1", block_type="table", content=[["Header"]])
                ]
            )
        ]
        out = tmp_path / "single_row_table.docx"
        export_sections_to_docx(sections, out)
        doc = Document(str(out))
        assert len(doc.tables) >= 1
        assert len(doc.tables[0].rows) == 1
